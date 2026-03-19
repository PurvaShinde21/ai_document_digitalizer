"""
OCR Service — Full pipeline running inside the FastAPI process.

Flow:
  file path → load images → preprocess → HuggingFace Inference API
           → postprocess text → generate TXT / DOCX / PDF outputs

No local model download. Uses HuggingFace Inference API (free tier:
30,000 req/month) with microsoft/trocr-small-handwritten (~270 MB vs
1.3 GB for the base variant).
"""

import io
import os
import re

import httpx
from PIL import Image

from config import settings
from models.job_model import JobStatus
from services.queue_service import update_job_status


# ── HuggingFace Inference API ────────────────────────────────────────────────

def _call_hf_api(image: Image.Image) -> str:
    """Send a PIL image to the HuggingFace Inference API, return predicted text."""
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    buf.seek(0)

    headers = {"Authorization": f"Bearer {settings.HF_API_TOKEN}"}
    response = httpx.post(
        settings.hf_api_url,
        headers=headers,
        content=buf.read(),
        timeout=60.0,
    )
    response.raise_for_status()
    result = response.json()

    if isinstance(result, list) and result:
        return result[0].get("generated_text", "")
    if isinstance(result, dict) and "error" in result:
        raise RuntimeError(f"HuggingFace API error: {result['error']}")
    return ""


# ── Image Preprocessing ──────────────────────────────────────────────────────

def _preprocess(image: Image.Image) -> Image.Image:
    """Convert to RGB and scale to a height TrOCR handles well."""
    if image.mode != "RGB":
        image = image.convert("RGB")
    w, h = image.size
    target_h = 384
    if h != target_h:
        ratio = target_h / h
        image = image.resize((int(w * ratio), target_h), Image.LANCZOS)
    return image


def _load_images(file_path: str) -> list[Image.Image]:
    """Return a list of PIL images from an image file or PDF."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        from pdf2image import convert_from_path
        return convert_from_path(file_path, dpi=200)
    return [Image.open(file_path)]


# ── Text Postprocessing ──────────────────────────────────────────────────────

def _postprocess(text: str) -> str:
    """Clean common OCR artefacts from the predicted text."""
    text = re.sub(r" {2,}", " ", text)          # collapse multiple spaces
    text = re.sub(r"\n{3,}", "\n\n", text)       # max two consecutive newlines
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text) # rejoin hyphenated line-breaks
    return text.strip()


# ── Output Generators ────────────────────────────────────────────────────────

def _save_txt(text: str, out_dir: str) -> str:
    path = os.path.join(out_dir, "result.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path


def _save_docx(text: str, out_dir: str) -> str:
    from docx import Document

    doc = Document()
    doc.add_heading("Digitized Document", level=1)
    for para in text.split("\n\n"):
        doc.add_paragraph(para.strip())
    path = os.path.join(out_dir, "result.docx")
    doc.save(path)
    return path


def _save_pdf(text: str, out_dir: str) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    path = os.path.join(out_dir, "result.pdf")
    doc_pdf = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Digitized Document", styles["Title"]),
        Spacer(1, 12),
    ]
    for para in text.split("\n\n"):
        clean = para.strip().replace("\n", "<br/>")
        story.append(Paragraph(clean, styles["Normal"]))
        story.append(Spacer(1, 8))
    doc_pdf.build(story)
    return path


# ── Main Pipeline (called as FastAPI BackgroundTask) ─────────────────────────

def run_ocr_pipeline(job_id: str, file_path: str, out_dir: str) -> None:
    """
    Full OCR pipeline orchestration.
    Runs in a background thread via FastAPI BackgroundTasks.
    Updates job status in Redis at each stage.
    """
    try:
        update_job_status(job_id, JobStatus.PROCESSING)

        # 1. Load images (handles both images and multi-page PDFs)
        images = _load_images(file_path)

        # 2. OCR each page via HuggingFace Inference API
        page_texts: list[str] = []
        for img in images:
            img = _preprocess(img)
            text = _call_hf_api(img)
            page_texts.append(text)

        # 3. Merge pages and clean up
        full_text = "\n\n--- Page Break ---\n\n".join(page_texts)
        full_text = _postprocess(full_text)

        # 4. Generate output files
        _save_txt(full_text, out_dir)
        _save_docx(full_text, out_dir)
        _save_pdf(full_text, out_dir)

        update_job_status(
            job_id,
            JobStatus.COMPLETED,
            output_formats=["txt", "docx", "pdf"],
        )

    except Exception as exc:
        update_job_status(job_id, JobStatus.FAILED, error=str(exc))
