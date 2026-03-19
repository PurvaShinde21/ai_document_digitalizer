"""Generate a formatted Word (.docx) document from OCR results."""

import os

from docx import Document
from docx.shared import Pt


def generate_docx(text: str, output_dir: str, filename: str = "result.docx") -> str:
    """
    Create a .docx file with the OCR text.
    Each double-newline-separated block becomes a paragraph.
    Returns the file path.
    """
    doc = Document()

    # Title
    title = doc.add_heading("Digitized Document", level=1)
    title.runs[0].font.size = Pt(18)

    # Body — split on paragraph breaks
    for paragraph in text.split("\n\n"):
        para = paragraph.strip()
        if para:
            doc.add_paragraph(para)

    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path
